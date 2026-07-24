"""Generate a .docx renovation proposal from a Quotation."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..calculators.quotation import Quotation
from ..calculators.rental_yield import RentalYieldCalculator


def _money(value: float) -> str:
    return f"NT$ {value:,.0f}"


def generate_word(
    quote: Quotation,
    output_path: str | Path,
    rental: RentalYieldCalculator | None = None,
) -> Path:
    """Render ``quote`` to a Word document at ``output_path``."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    title = doc.add_heading("Renovation Proposal", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(quote.project_name)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)

    diff = quote.difficulty.breakdown()
    doc.add_paragraph(
        f"Difficulty multiplier: {diff['multiplier']:.2f}x   "
        f"(floor +{diff['floor_factor']:.2f}, age +{diff['age_factor']:.2f}, "
        f"condition +{diff['condition_factor']:.2f})"
    )

    # -- per-category tables ----------------------------------------------
    for cat in quote.category_breakdown():
        doc.add_heading(cat["display_name"], level=1)
        doc.add_paragraph(cat["description"])

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Qty"
        hdr[2].text = "Unit price"
        hdr[3].text = "Amount"
        for item in cat["items"]:
            row = table.add_row().cells
            row[0].text = item["name"]
            row[1].text = f"{item['quantity']:g} {item['unit']}"
            row[2].text = _money(item["unit_price"])
            row[3].text = _money(item["base_cost"])

        doc.add_paragraph(
            f"Subtotal (x{cat['multiplier']:.2f}): {_money(cat['subtotal'])}"
        ).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # -- totals ------------------------------------------------------------
    doc.add_heading("Summary", level=1)
    summary = doc.add_table(rows=0, cols=2)
    summary.style = "Light List Accent 1"
    for label, value in (
        ("Trade subtotal", quote.trade_subtotal),
        (f"Supervision fee ({quote.supervision_rate:.0%})", quote.supervision_fee),
        ("Grand total", quote.grand_total),
    ):
        row = summary.add_row().cells
        row[0].text = label
        row[1].text = _money(value)

    # -- rental yield ------------------------------------------------------
    if rental is not None:
        doc.add_heading("Rental Yield Analysis", level=1)
        ry = rental.as_dict()
        analysis = doc.add_table(rows=0, cols=2)
        analysis.style = "Light List Accent 1"
        rows = (
            ("Expected monthly rent", _money(ry["monthly_rent"])),
            ("Effective monthly rent", _money(ry["effective_monthly_rent"])),
            ("Total investment", _money(ry["total_investment"])),
            ("Monthly net profit", _money(ry["monthly_net_profit"])),
            ("Net annual yield", f"{ry['net_yield']:.2%}"),
            (
                "Payback period",
                "n/a"
                if ry["payback_months"] == float("inf")
                else f"{ry['payback_months']:.1f} months",
            ),
        )
        for label, value in rows:
            cells = analysis.add_row().cells
            cells[0].text = label
            cells[1].text = value

    doc.add_paragraph()
    footer = doc.add_paragraph("Demo document — sample data only, not a binding quote.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    doc.save(output_path)
    return output_path
